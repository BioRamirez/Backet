#--------------## Cargar librerias necesarias------------------------------
# Si no las tienes instaladas, ejecuta esta celda una vez:
# Salir del interprete con: exit() exit() python   pip install tabulate pandas numpy scipy scikit-bio openpyxl
#
# !pip install pandas numpy matplotlib tabulate openpyxl

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from tabulate import tabulate
import openpyxl

#--------------## Leer archivo y revisar columnas------------------------------
# Ruta del archivo
ruta = r"D:\CORPONOR 2025\Backet\python_Proyect\data\POF_ZULIA_2025_BD_AVES_MAMIFEROS.xlsx"

# Leer el archivo Excel
Registros = pd.read_excel(ruta)

# Mostrar las primeras filas
print("📄 Primeras filas del archivo:")
print(Registros.head())

# Mostrar nombres de las columnas
print("\n📋 Columnas del DataFrame:")
print(Registros.columns)

#------------------------------## Crear matriz de abundancia------------------------------

#------------------------------Calcular indices de diversidad------------------------------
import re

def generar_abreviacion(nombre):
    """
    Genera abreviaciones automáticas a partir de nombres de coberturas.
    Ejemplo: 'Bosque de galería y ripario' → 'Bgr'
    """
    # Convertir a minúsculas y dividir en palabras
    palabras = nombre.lower().split()

    # Eliminar conectores comunes
    palabras = [p for p in palabras if p not in ['de', 'del', 'la', 'el', 'y', 'con', 'en', 'los', 'las']]

    # Tomar la primera letra de cada palabra
    abreviacion = ''.join([p[0] for p in palabras])

    # Asegurar que tenga al menos 3 caracteres (rellena si es corta)
    if len(abreviacion) < 3:
        abreviacion = abreviacion.ljust(3, '_')

    return abreviacion.capitalize()


def abreviar_coberturas(df, columna='COBERTURA'):
    """
    Crea un diccionario de abreviaciones y reemplaza los nombres en el DataFrame.
    """
    coberturas_unicas = df[columna].unique()
    abreviaciones = {c: generar_abreviacion(c) for c in coberturas_unicas}

    print("🔤 Abreviaciones generadas automáticamente:")
    for original, abrev in abreviaciones.items():
        print(f"  {original} → {abrev}")

    # Reemplazar en el DataFrame
    df[columna] = df[columna].replace(abreviaciones)

    return df, abreviaciones


# --- Aplicar las abreviaciones en el DataFrame ---
Registros, abreviaciones_cobertura = abreviar_coberturas(Registros, columna='COBERTURA')

print("✅ Abreviaciones aplicadas a la columna 'COBERTURA':")
print(Registros['COBERTURA'].unique())


import pandas as pd
import numpy as np
from scipy.stats import entropy
import matplotlib.pyplot as plt

# --- 1. Crear matriz de abundancia por cobertura ---
matriz_abundancia = Registros.pivot_table(
    index='COBERTURA',
    columns='ESPECIE',
    values='INDIVIDUOS',
    aggfunc='sum',
    fill_value=0
)

print("✅ Matriz de abundancia creada:")
print(matriz_abundancia.head())

# --- 2. Función para calcular índices ecológicos ---
def calcular_indices(abundancias):
    abundancias = np.array(abundancias)
    N = abundancias.sum()
    S = np.count_nonzero(abundancias)
    
    if N == 0 or S == 0:
        return {
            'Riqueza (S)': 0,
            'Abundancia (N)': 0,
            'Shannon (H\')': 0,
            'Simpson (1-D)': 0,
            'Dominancia (D)': 0,
            'Equidad (J\')': 0,
            'Margalef (DMg)': 0,
            'Menhinick (DMn)': 0
        }

    # Proporciones
    p = abundancias / N

    # Índices de diversidad
    shannon = entropy(p, base=np.e)
    simpson = 1 - np.sum(p**2)
    dominancia = np.sum(p**2)
    pielou = shannon / np.log(S)
    margalef = (S - 1) / np.log(N)
    menhinick = S / np.sqrt(N)

    return {
        'Riqueza (S)': S,
        'Abundancia (N)': N,
        'Shannon (H\')': shannon,
        'Simpson (1-D)': simpson,
        'Dominancia (D)': dominancia,
        'Equidad (J\')': pielou,
        'Margalef (DMg)': margalef,
        'Menhinick (DMn)': menhinick
    }

# --- 3. Calcular índices por cobertura ---
indices_diversidad = matriz_abundancia.apply(calcular_indices, axis=1, result_type='expand')

print("\n📊 Índices de diversidad por cobertura:")
print(indices_diversidad)


# --- 4. Calcular índices totales (todas las coberturas combinadas) ---
abundancia_total = matriz_abundancia.sum(axis=0)
indices_totales = calcular_indices(abundancia_total)

print("\n🌍 Índices de diversidad total (todas las coberturas combinadas):")
for k, v in indices_totales.items():
    print(f"{k}: {v:.4f}")

# Convertir los índices totales a DataFrame
indices_totales_df = pd.DataFrame([indices_totales], index=['Total'])


# Combinar ambos DataFrames
indices_combinados = pd.concat([indices_diversidad, indices_totales_df])

print("📊 Tabla combinada de índices de diversidad:")
print(indices_combinados)

indices_combinados = indices_combinados.round(4)
print("📊 Tabla combinada de índices de diversidad (redondeada):")
print(indices_combinados)

indices_pivot = indices_combinados.T
indices_pivot.index.name = 'Índice'
indices_pivot.columns.name = 'Cobertura'

print(indices_pivot)



indices_pivot.to_excel("D:/CORPONOR 2025/Backet/python_Proyect/Resultados/Indices_Diversidad_Combinados.xlsx")
# Confirmar la ubicación del archivo guardado
print(f"✅ Archivo exportado correctamente en:\nD:/CORPONOR 2025/Backet/python_Proyect/Resultados/Indices_Diversidad_Combinados.xlsx")

#--------------------Formatear grafica--------------------

#---------------------------------- Reparar y formatear archivo de Indices_Diversidad_Combinados -----------------------
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
import os

# --- Rutas ---
ruta_original = r"D:\CORPONOR 2025\Backet\python_Proyect\Resultados\Indices_Diversidad_Combinados.xlsx"
ruta_limpia = r"D:\CORPONOR 2025\Backet\python_Proyect\Resultados\Indices_Diversidad_Combinados.xlsx"

# --- Verificar existencia ---
if not os.path.exists(ruta_original):
    raise FileNotFoundError(f"⚠️ No se encontró el archivo: {ruta_original}")

# --- Leer archivo dañado con pandas ---
try:
    df = pd.read_excel(ruta_original)
    print("✅ Archivo leído correctamente con pandas.")
except Exception as e:
    raise RuntimeError(f"❌ No se pudo leer el archivo: {e}")

# --- Reescribir el archivo limpio ---
df.to_excel(ruta_limpia, index=False)
print(f"🧹 Archivo reparado y guardado como:\n{ruta_limpia}")

# --- Aplicar formato con openpyxl ---
from openpyxl import load_workbook

wb = load_workbook(ruta_limpia)
ws = wb.active

# --- Estilos base ---
header_fill = PatternFill(start_color='BFD8B8', end_color='BFD8B8', fill_type='solid')
header_font = Font(bold=True, color='000000', name='Calibri')
center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
thin_border = Border(
    left=Side(style='thin', color='000000'),
    right=Side(style='thin', color='000000'),
    top=Side(style='thin', color='000000'),
    bottom=Side(style='thin', color='000000')
)

# --- Aplicar formato y reemplazar vacíos ---
for row in ws.iter_rows():
    for cell in row:
        if cell.value is None or str(cell.value).strip() == '':
            cell.value = '-'
        cell.alignment = center_align
        cell.border = thin_border

# --- Encabezado ---
for cell in ws[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = center_align

# --- Ajustar ancho de columnas ---
for col in ws.columns:
    max_length = 0
    column = get_column_letter(col[0].column)
    for cell in col:
        if cell.value:
            length = len(str(cell.value))
            if length > max_length:
                max_length = length
    ws.column_dimensions[column].width = max_length + 3

# --- Ajustar altura de filas ---
for row in ws.iter_rows():
    ws.row_dimensions[row[0].row].height = 18

# --- Guardar cambios ---
wb.save(ruta_limpia)
print(f'📘 Archivo formateado y reparado correctamente:\n{ruta_limpia}')

#-------------Fin del formateo del archivo----------------------------
#---------------------------Interpreetar tabla de Indices de diversidad---------------------------
import pandas as pd
import numpy as np

def clasificar_indice(valor, rangos):
    """Clasifica un valor numérico según los rangos definidos."""
    for categoria, (minv, maxv) in rangos.items():
        if minv <= valor < maxv:
            return categoria
    return "Fuera de rango"

def interpretar_indices_completo(df):
    """
    Interpreta automáticamente una tabla de índices ecológicos.
    Incluye descripción de cada índice, su interpretación y nivel (bajo, medio, alto).
    Filas = índices, columnas = coberturas (la última puede ser 'Total').
    """
    interpretaciones = []
    
    # --- Rangos ecológicos generales ---
    rangos_dict = {
        "Shannon": {"Baja": (0, 2), "Media": (2, 3.5), "Alta": (3.5, 100)},
        "Simpson": {"Baja": (0, 0.8), "Media": (0.8, 0.95), "Alta": (0.95, 1.01)},
        "Dominancia": {"Alta": (0.1, 1), "Media": (0.05, 0.1), "Baja": (0, 0.05)},
        "Equidad": {"Baja": (0, 0.6), "Media": (0.6, 0.8), "Alta": (0.8, 1.01)},
        "Margalef": {"Baja": (0, 5), "Media": (5, 15), "Alta": (15, 100)},
        "Menhinick": {"Baja": (0, 2), "Media": (2, 4), "Alta": (4, 100)}
    }

    # --- Descripciones breves de cada índice ---
    descripciones = {
        "Riqueza": "El índice de **riqueza (S)** representa el número total de especies registradas en una cobertura. No considera la abundancia, solo cuántas especies hay.",
        "Abundancia": "La **abundancia (N)** refleja el número total de individuos registrados; una alta abundancia puede indicar hábitats más productivos o mejor muestreados.",
        "Shannon": "El índice de **Shannon (H’)** mide la diversidad teniendo en cuenta tanto la riqueza de especies como su equidad. Valores altos indican comunidades más diversas y equilibradas.",
        "Simpson": "El índice de **Simpson (1–D)** expresa la probabilidad de que dos individuos seleccionados al azar pertenezcan a especies diferentes. Valores cercanos a 1 reflejan alta diversidad.",
        "Dominancia": "El índice de **Dominancia (D)** mide el grado en que una o pocas especies dominan el ensamblaje. Valores altos indican dominancia de pocas especies.",
        "Equidad": "La **equidad (J’)** describe cuán uniformemente se distribuyen los individuos entre las especies. Valores altos indican distribución equitativa.",
        "Margalef": "El índice de **Margalef (DMg)** ajusta la riqueza de especies en función del número de individuos, útil para comparar entre coberturas con distinto esfuerzo de muestreo.",
        "Menhinick": "El índice de **Menhinick (DMn)** también ajusta la riqueza según la abundancia total, proporcionando una medida estandarizada de riqueza relativa."
    }

    # --- Copia de seguridad del DataFrame ---
    datos = df.copy()
    
    # Identificar la columna de Totales si existe
    col_total = None
    for col in datos.columns:
        if "total" in col.lower():
            col_total = col
            break
    
    coberturas = [c for c in datos.columns if c != col_total]

    # --- Interpretación por índice ---
    for indice in datos.index:
        nombre_limpio = indice.split(" ")[0].replace("(", "").replace(")", "")
        valores = datos.loc[indice, coberturas]

        if np.issubdtype(valores.dtype, np.number):
            cobertura_max = valores.idxmax()
            cobertura_min = valores.idxmin()
            val_max = valores.max()
            val_min = valores.min()

            # --- Descripción ---
            tipo = next((k for k in rangos_dict.keys() if k.lower() in nombre_limpio.lower()), None)
            desc = next((v for k, v in descripciones.items() if k.lower() in nombre_limpio.lower()), None)
            if not desc:
                desc = f"El índice **{indice}** evalúa un aspecto ecológico particular del ensamblaje de especies."

            interpretaciones.append(f"\n📘 {desc}")

            # --- Clasificación ---
            if tipo:
                nivel_max = clasificar_indice(val_max, rangos_dict[tipo])
                nivel_min = clasificar_indice(val_min, rangos_dict[tipo])
                interpretaciones.append(
                    f"🔹 En este índice, la cobertura con valor más alto es **{cobertura_max}** "
                    f"({val_max:.3f}, categoría {nivel_max.lower()}) y la más baja es **{cobertura_min}** "
                    f"({val_min:.3f}, categoría {nivel_min.lower()})."
                )
            else:
                interpretaciones.append(
                    f"🔹 La cobertura con mayor valor de **{indice}** es **{cobertura_max}** "
                    f"({val_max:.3f}), mientras que la más baja es **{cobertura_min}** ({val_min:.3f})."
                )

    # --- Interpretación general ---
    interpretaciones.append("\n🧩 **Síntesis ecológica general:**")
    interpretaciones.append(
        "Altos valores en los índices de **Shannon** y **Simpson** reflejan comunidades con gran diversidad y "
        "una distribución equilibrada de individuos entre especies. "
        "En contraste, valores altos de **Dominancia** indican concentración de abundancia en pocas especies. "
        "Los índices de **Equidad** expresan el grado de uniformidad en la distribución de individuos, "
        "mientras que **Margalef** y **Menhinick** complementan la evaluación de la riqueza relativa ajustada por abundancia."
    )

    if col_total:
        interpretaciones.append(
            f"\n📊 Finalmente, la columna **{col_total}** resume los valores combinados del conjunto total de coberturas, "
            "brindando una visión general de la diversidad del ecosistema muestreado."
        )

    return "\n".join(interpretaciones)


# --- Ejemplo de uso ---
interpretacion = interpretar_indices_completo(indices_pivot)
print(interpretacion)

#----------------------------Crear informe W0rd----------------------------
#pip install python-docx
from docx import Document
from docx.shared import Inches

# --- 2️⃣ Crear documento Word ---
ruta_salida = "D:/CORPONOR 2025/Backet/python_Proyect/Resultados/Interpretacion_Indices_Diversidad.docx"
doc = Document()

# --- Título principal ---
doc.add_heading("Informe de Índices de Diversidad", level=1)

# --- Subtítulo ---
doc.add_heading("Tabla de Índices Calculados", level=2)

# --- 3️⃣ Insertar la tabla en el documento ---
tabla = doc.add_table(rows=1, cols=len(indices_pivot.columns) + 1)
tabla.style = 'Table Grid'

# --- Encabezados ---
hdr_cells = tabla.rows[0].cells
hdr_cells[0].text = "Índice"
for i, col in enumerate(indices_pivot.columns):
    hdr_cells[i + 1].text = col

# --- Filas con datos ---
for idx, row in indices_pivot.iterrows():
    row_cells = tabla.add_row().cells
    row_cells[0].text = str(idx)
    for j, valor in enumerate(row):
        row_cells[j + 1].text = f"{valor:.4f}" if isinstance(valor, (float, int)) else str(valor)

# --- 4️⃣ Espacio y título para la interpretación ---
doc.add_paragraph("\n")
doc.add_heading("Interpretación Automática de los Resultados", level=2)
doc.add_paragraph(interpretacion)

# --- 5️⃣ Guardar archivo ---
doc.save(ruta_salida)
print(f"✅ Archivo con tabla e interpretación guardado en:\n{ruta_salida}")


#----------------------------Fin del codigo----------------------------




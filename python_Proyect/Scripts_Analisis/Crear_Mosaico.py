

import os
import math
import pandas as pd
from PIL import Image, ImageDraw, ImageFont

# ========== CONFIGURACIÓN ==========
#"D:\Forestal Consultores\2025\Forestal_2025\Informacion Primaria\FOTOS\COPIAS_IMAGENES_RECUPERADAS\AVES_INFORME"

carpeta_fotos = r"D:\Forestal Consultores\2026\FAUNA\BD\ANFIBIOS\FotosInformeAnfCopia"
archivo_excel = r"D:\Forestal Consultores\2026\FAUNA\BD\BD_SANROQUE.xlsx"
col_orden = "Orden"
col_especie = "ESPECIE"

ext_validas = {".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG"}

# Carpeta de salida EXACTA: python_Proyect/Resultados/Mosaico
base_proyect = os.path.dirname(os.path.dirname(archivo_excel))  # Subir desde /data a /python_Proyect
carpeta_salida = os.path.join(base_proyect, "Resultados", "Mosaico")
os.makedirs(carpeta_salida, exist_ok=True)


# Miniaturas
THUMB_W = 480
THUMB_H = 360

# ========== CARGAR BASE ==========
df = pd.read_excel(archivo_excel)
df[col_especie] = df[col_especie].astype(str).str.strip()
df[col_orden] = df[col_orden].astype(str).str.strip()
ordenes = df[col_orden].unique()

# ========== BUSCAR FOTOS ==========
def buscar_fotos_especie(nombre_especie):
    fotos = []
    nl = nombre_especie.lower()
    for archivo in os.listdir(carpeta_fotos):
        if any(archivo.endswith(ext) for ext in ext_validas):
            if nl in archivo.lower():
                fotos.append(os.path.join(carpeta_fotos, archivo))
    return fotos

# ========== DIMENSIONES ADAPTATIVAS ==========
def dimensiones_mosaico(n):
    if n == 1:
        return None
    if n == 2:
        return (2, 1)
    if 3 <= n <= 6:
        cols = 2
        rows = math.ceil(n / cols)
        return (cols, rows)
    if 7 <= n <= 10:
        cols = 3
        rows = math.ceil(n / cols)
        return (cols, rows)
    cols = 4
    rows = math.ceil(n / cols)
    return (cols, rows)

# ========== UTIL: obtener bbox real de "ink" para un texto ==========
def ink_bbox_for_text(text, font):
    """
    Renderiza el texto en una máscara y devuelve el bbox real de píxeles pintados
    y la imagen máscara (L) junto al tamaño usado.
    """
    # Primero obtener una estimación de tamaño (textbbox es útil como punto de partida)
    dummy = Image.new("L", (1, 1), 0)
    dd = ImageDraw.Draw(dummy)
    bbox = dd.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]

    # Creamos una máscara más amplia para evitar recorte
    pad = 32
    W = tw + pad * 2
    H = th + pad * 2
    mask = Image.new("L", (W, H), 0)
    md = ImageDraw.Draw(mask)
    # Pintamos el texto centrado en la máscara (colocamos en pad,pad, no hace falta más)
    md.text((pad, pad), text, font=font, fill=255)

    # getbbox() sobre la máscara devuelve el bbox exacto de pixeles no negros (ink)
    ink_bbox = mask.getbbox()  # (left, upper, right, lower) o None si está vacío
    return ink_bbox, mask

# ========== CREAR MOSAICO (CORREGIDO Y CENTRADO POR "INK") ==========
def crear_mosaico_from_list(pares_especie_ruta, salida_mosaico, salida_txt, titulo):
    n = len(pares_especie_ruta)
    dims = dimensiones_mosaico(n)
    if dims is None:
        return False

    cols, rows = dims
    ancho_total = cols * THUMB_W
    alto_total = rows * THUMB_H

    mosaico = Image.new("RGB", (ancho_total, alto_total), (255, 255, 255))

    # Letras generadas
    from string import ascii_lowercase
    letras = []
    if n <= 26:
        letras = list(ascii_lowercase[:n])
    else:
        letras = list(ascii_lowercase)
        extra = 1
        i = 26
        while i < n:
            for c in ascii_lowercase:
                if i >= n:
                    break
                letras.append(f"{c}{extra}")
                i += 1
            extra += 1

    # Fuente: intenta Arial, si no existe usa la default
    try:
        font = ImageFont.truetype("arial.ttf", 44)
    except:
        font = ImageFont.load_default()

    posiciones_letras = []
    lista_txt = []

    # === 1. PEGAR THUMBNAILS SIN LETRA ===
    for idx, (especie, ruta) in enumerate(pares_especie_ruta):
        try:
            img = Image.open(ruta).convert("RGB")
            img = img.resize((THUMB_W, THUMB_H))
        except Exception as e:
            print("⚠ Error abriendo:", ruta, e)
            continue

        c = idx % cols
        r = idx // cols
        x = c * THUMB_W
        y = r * THUMB_H
        mosaico.paste(img, (x, y))

        posiciones_letras.append((idx, x, y))
        lista_txt.append(f"{letras[idx]}. {especie}")
        

    parrafo = ", ".join(lista_txt)
    import textwrap
    parrafo = textwrap.fill(parrafo, width=90)


    # === 2. DIBUJAR LETRAS Y RECUADROS (CENTRADO REAL USANDO MÁSCARA) ===
    draw = ImageDraw.Draw(mosaico)

    for (idx, x0, y0) in posiciones_letras:
        letra = letras[idx]

        # Obtenemos el bbox "ink" real y la máscara renderizada
        ink_bbox, mask = ink_bbox_for_text(letra, font)
        if ink_bbox is None:
            # texto vacío por alguna razón: fallback a textbbox
            bbox = draw.textbbox((0, 0), letra, font=font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            ink_left, ink_top, ink_right, ink_bottom = 0, 0, tw, th
            mask = None
        else:
            ink_left, ink_top, ink_right, ink_bottom = ink_bbox
            tw = ink_right - ink_left
            th = ink_bottom - ink_top

        # Padding para el cuadro
        pad_x = 14
        pad_y = 12

        # Esquina inferior derecha del thumbnail (separación del borde)
        gap = 16
        corner_x = x0 + THUMB_W - gap
        corner_y = y0 + THUMB_H - gap

        # Dimensiones del cuadro negro
        cuadro_w = tw + pad_x * 2
        cuadro_h = th + pad_y * 2

        # Coordenadas del cuadro negro (colocado en la esquina inferior derecha)
        X1 = corner_x - cuadro_w
        Y1 = corner_y - cuadro_h
        X2 = corner_x
        Y2 = corner_y

        # Dibujar rectángulo negro
        draw.rectangle([X1, Y1, X2, Y2], fill=(0, 0, 0))

        # Calcular posición exacta donde pegar la "ink mask" para que quede centrada
        # text_x/text_y son coordenadas en el mosaico donde debe empezar la máscara
        text_x = X1 + (cuadro_w - tw) // 2
        text_y = Y1 + (cuadro_h - th) // 2

        # Si tenemos la máscara con padding (pad en la función), debemos desplazar para
        # que el recorte de ink_bbox coincida.
        if mask is not None and ink_bbox is not None:
            # la máscara tiene pad interno; queremos pegar solo la región ink_bbox
            # dentro de la máscara en (text_x, text_y)
            ink_crop = mask.crop((ink_left, ink_top, ink_right, ink_bottom))  # L image (blanco sobre negro)
            # Crear una imagen RGBA para pegar el texto blanco con transparencia
            text_rgba = Image.new("RGBA", (tw, th), (255, 255, 255, 0))
            # Poner blancos donde mask tiene tinta (usamos mask como alpha)
            white = Image.new("RGBA", (tw, th), (255, 255, 255, 255))
            text_rgba.paste(white, (0, 0), mask=ink_crop)
            # Pegar text_rgba sobre el mosaico en text_x,text_y
            mosaico.paste(text_rgba, (int(text_x), int(text_y)), text_rgba)
        else:
            # Fallback: usar draw.text centrado por tamaño aproximado
            draw.text((text_x, text_y), letra, fill=(255, 255, 255), font=font)

    # === GUARDAR ===
    mosaico.save(salida_mosaico, "JPEG", quality=95)

    with open(salida_txt, "w", encoding="utf-8") as f:
        f.write(f"{titulo}\n\n")
        f.write(parrafo)

    print("✔ Guardado:", salida_mosaico)
    print("✔ TXT generado:", parrafo)
    return True

# ========== PROCESO PRINCIPAL ==========
mosaico_unafoto = []

for orden in ordenes:
    print(f"\n=== Procesando orden: {orden} ===")
    especies = df[df[col_orden] == orden][col_especie].unique()

    pares = []
    for especie in especies:
        rutas = buscar_fotos_especie(especie)
        for r in rutas:
            pares.append((especie, r))

    if len(pares) == 0:
        continue

    if len(pares) == 1:
        especie, ruta = pares[0]
        mosaico_unafoto.append((f"{especie} ({orden})", ruta))
        continue

    salida_mosaico = os.path.join(carpeta_salida, f"MOSAICO_{orden}.jpg")
    salida_txt = os.path.join(carpeta_salida, f"{orden}_mosaico.txt")
    titulo = f"Figura — Orden: {orden} (Total fotos: {len(pares)})"

    crear_mosaico_from_list(pares, salida_mosaico, salida_txt, titulo)

# ========== MOSAICO MULTI-ORDEN ==========
if mosaico_unafoto:
    print("\n=== Creando MOSAICO MULTI-ORDEN ===")
    salida_mosaico = os.path.join(carpeta_salida, "MOSAICO_MULTI_ORDEN.jpg")
    salida_txt = os.path.join(carpeta_salida, "MULTI_ORDEN_mosaico.txt")
    titulo = f"Figura — Órdenes con una sola foto (Total: {len(mosaico_unafoto)})"

    crear_mosaico_from_list(mosaico_unafoto, salida_mosaico, salida_txt, titulo)

print("\n🎉 PROCESO COMPLETO. Archivos en:", carpeta_salida)



























































































import os
import math
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========= CONFIGURACIÓN ==========

carpeta_fotos = r"D:\Forestal Consultores\2026\FAUNA\BD\ANFIBIOS\FotosInformeAnfCopia"
archivo_excel = r"D:\Forestal Consultores\2026\FAUNA\BD\BD_SANROQUE.xlsx"

col_orden = "Orden"
col_especie = "ESPECIE"
col_familia = "Familia"

ext_validas = {".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG"}

# Carpeta de salida
base_proyect = os.path.dirname(os.path.dirname(archivo_excel))
carpeta_salida = os.path.join(base_proyect, "Resultados", "Mosaico_Word")
os.makedirs(carpeta_salida, exist_ok=True)

salida_word = os.path.join(carpeta_salida, "MOSAICOS_ORDENES.docx")

# ========= CARGAR DATOS ==========
df = pd.read_excel(archivo_excel)
df[col_especie] = df[col_especie].astype(str).str.strip()
df[col_orden] = df[col_orden].astype(str).str.strip()
df[col_familia] = df[col_familia].astype(str).str.strip()

ordenes = df[col_orden].unique()


# ========= FUNCIÓN: BUSCAR FOTOS ==========
def buscar_fotos_especie(nombre_especie):
    fotos = []
    nl = nombre_especie.lower()
    for archivo in os.listdir(carpeta_fotos):
        if any(archivo.endswith(ext) for ext in ext_validas):
            if nl in archivo.lower():
                fotos.append(os.path.join(carpeta_fotos, archivo))
    return fotos


# ========= CREAR DOCUMENTO ==========
doc = Document()
doc.add_heading("Mosaicos por Orden – Fauna", level=1)

imagenes_usadas = set()
especies_unafoto = []

# ========= PROCESAR POR ORDEN ==========
for orden in ordenes:

    especies = df[df[col_orden] == orden].drop_duplicates(subset=[col_especie])
    pares = []  # (especie, familia, ruta)

    for _, row in especies.iterrows():
        esp = row[col_especie]
        fam = row[col_familia]
        rutas = buscar_fotos_especie(esp)

        for r in rutas:
            if r not in imagenes_usadas:
                pares.append((esp, fam, r))
                imagenes_usadas.add(r)

    if len(pares) == 0:
        continue

    if len(pares) == 1:
        esp, fam, ruta = pares[0]
        especies_unafoto.append((orden, esp, fam, ruta))
        continue

    # ---- Crear tabla ----
    doc.add_heading(f"Orden: {orden}", level=2)

    filas = math.ceil(len(pares) / 2)
    tabla = doc.add_table(rows=filas, cols=2)
    tabla.autofit = False

    idx = 0
    for r in range(filas):
        for c in range(2):
            if idx >= len(pares):
                break

            esp, fam, ruta = pares[idx]
            cell = tabla.rows[r].cells[c]

            # Eliminar contenido automático
            cell._element.clear()

            # Imagen centrada
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(ruta, width=Cm(7.33), height=Cm(5.5))
            except:
                p_img.add_run("[Error cargando imagen]").font.name = "Arial"

            # Texto centrado
            p_txt = cell.add_paragraph()
            p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_esp = p_txt.add_run(esp)
            run_esp.italic = True
            run_esp.font.name = "Arial"

            run_fam = p_txt.add_run(f" ({fam})")
            run_fam.font.name = "Arial"

            idx += 1

    doc.add_page_break()


# ========= TABLA FINAL MULTI-ORDEN ==========
if especies_unafoto:
    doc.add_heading("Órdenes con una sola foto", level=2)

    filas = math.ceil(len(especies_unafoto) / 2)
    tabla = doc.add_table(rows=filas, cols=2)

    idx = 0
    for r in range(filas):
        for c in range(2):
            if idx >= len(especies_unafoto):
                break

            orden, esp, fam, ruta = especies_unafoto[idx]
            cell = tabla.rows[r].cells[c]

            # Limpiar contenido
            cell._element.clear()

            # Imagen centrada
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(ruta, width=Cm(7.33), height=Cm(5.5))
            except:
                p_img.add_run("[Error cargando imagen]").font.name = "Arial"

            # Texto centrado
            p_txt = cell.add_paragraph()
            p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_esp = p_txt.add_run(esp)
            run_esp.italic = True
            run_esp.font.name = "Arial"

            run_info = p_txt.add_run(f" ({fam}) – Orden: {orden}")
            run_info.font.name = "Arial"

            idx += 1

# ========= GUARDAR ==========
doc.save(salida_word)
print("\n✔ Word generado en:", salida_word)






































































#-------CUANDO LAS IMAGENES SEAN MUY PESADAS-----------------

import os
import math
import pandas as pd
from PIL import Image
import tempfile
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========= CONFIGURACIÓN ==========

carpeta_fotos = r"D:\Forestal Consultores\2026\FAUNA\BD\ANFIBIOS\FotosInformeAnfCopia"
archivo_excel = r"D:\Forestal Consultores\2026\FAUNA\BD\BD_SANROQUE.xlsx"

col_orden = "Orden"
col_especie = "ESPECIE"
col_familia = "Familia"

ext_validas = {".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG", ".webp", ".tif", ".tiff"}

# Carpeta de salida
base_proyect = os.path.dirname(os.path.dirname(archivo_excel))
carpeta_salida = os.path.join(base_proyect, "Resultados", "Mosaico_Word")
os.makedirs(carpeta_salida, exist_ok=True)

salida_word = os.path.join(carpeta_salida, "MOSAICOS_ORDENES.docx")

# ========= FUNCIÓN ROBUSTA PARA INSERTAR IMÁGENES ==========
def insertar_imagen(parrafo, ruta_img, width_cm=7.33, height_cm=5.5,
                     peso_umbral_kb=1000, max_dim_umbral_px=3000):
    """
    Inserta una imagen en el párrafo garantizando compatibilidad:
    - Convierte transparencias a fondo blanco.
    - Convierte a RGB si es necesario.
    - Reduce resolución si es muy pesada o muy grande.
    - Guarda copia temporal en JPEG y la inserta.
    """
    try:
        if not os.path.isfile(ruta_img):
            parrafo.add_run(f"[Imagen no encontrada: {os.path.basename(ruta_img)}]").font.name = "Arial"
            return

        peso_kb = os.path.getsize(ruta_img) / 1024.0

        # Abrir imagen con PIL
        img = Image.open(ruta_img)

        # Manejar transparencia (RGBA, LA) o paletas con transparencia (modo 'P' + info['transparency'])
        needs_alpha_handling = False
        if img.mode in ("RGBA", "LA"):
            needs_alpha_handling = True
        elif img.mode == "P" and ("transparency" in img.info):
            needs_alpha_handling = True

        if needs_alpha_handling:
            # Convertir a RGBA para obtener canal alpha
            rgba = img.convert("RGBA")
            alpha = rgba.split()[-1]
            # Crear fondo blanco y pegar la imagen usando la máscara alpha
            bg = Image.new("RGB", rgba.size, (255, 255, 255))
            bg.paste(rgba, mask=alpha)
            img = bg  # ahora es RGB sin transparencia
        else:
            # Si no es RGB, convertir simple a RGB (evita errores con "P", "L", etc.)
            if img.mode != "RGB":
                img = img.convert("RGB")

        # Reducir resolución si la imagen es muy pesada o demasiado grande en dimensión
        if peso_kb > peso_umbral_kb or max(img.size) > max_dim_umbral_px:
            # thumbnail mantiene proporción
            img.thumbnail((2000, 2000))

        # Guardar copia temporal como JPEG para asegurar compatibilidad con python-docx/Word
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
        temp_path = temp.name
        try:
            img.save(temp_path, format="JPEG", quality=85)
        finally:
            temp.close()

        # Insertar la imagen temporal en el documento
        parrafo.add_run().add_picture(temp_path, width=Cm(width_cm), height=Cm(height_cm))

        # opcional: se podría eliminar el archivo temporal aquí si se desea:
        # os.remove(temp_path)

    except Exception as e:
        # Mensaje claro en el documento para identificar la especie/archivo problemático
        parrafo.add_run(f"[Error cargando imagen: {e}]").font.name = "Arial"


# ========= CARGAR DATOS ==========
df = pd.read_excel(archivo_excel)
df[col_especie] = df[col_especie].astype(str).str.strip()
df[col_orden] = df[col_orden].astype(str).str.strip()
df[col_familia] = df[col_familia].astype(str).str.strip()

ordenes = df[col_orden].unique()


# ========= FUNCIÓN: BUSCAR FOTOS ==========
def buscar_fotos_especie(nombre_especie):
    fotos = []
    nl = nombre_especie.lower()

    # proteger contra carpeta inexistente
    if not os.path.isdir(carpeta_fotos):
        return fotos

    for archivo in os.listdir(carpeta_fotos):
        if any(archivo.endswith(ext) for ext in ext_validas):
            if nl in archivo.lower():
                fotos.append(os.path.join(carpeta_fotos, archivo))

    return fotos


# ========= CREAR DOCUMENTO ==========
doc = Document()
doc.add_heading("Mosaicos por Orden – Fauna", level=1)

imagenes_usadas = set()
especies_unafoto = []

# ========= PROCESAR POR ORDEN ==========
for orden in ordenes:

    especies = df[df[col_orden] == orden].drop_duplicates(subset=[col_especie])
    pares = []  # (especie, familia, ruta)

    for _, row in especies.iterrows():
        esp = row[col_especie]
        fam = row[col_familia]
        rutas = buscar_fotos_especie(esp)

        for r in rutas:
            if r not in imagenes_usadas:
                pares.append((esp, fam, r))
                imagenes_usadas.add(r)

    if len(pares) == 0:
        continue

    if len(pares) == 1:
        especies_unafoto.append((orden, pares[0][0], pares[0][1], pares[0][2]))
        continue
    pares.sort(key=lambda x: x[1])

    # ---- Crear tabla ----
    doc.add_heading(f"Orden: {orden}", level=2)

    filas = math.ceil(len(pares) / 2)
    tabla = doc.add_table(rows=filas, cols=2)
    tabla.autofit = False

    idx = 0
    for r in range(filas):
        for c in range(2):
            if idx >= len(pares):
                break

            esp, fam, ruta = pares[idx]
            cell = tabla.rows[r].cells[c]

            cell._element.clear()

            # Imagen centrada
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            insertar_imagen(p_img, ruta)

            # Texto
            p_txt = cell.add_paragraph()
            p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_esp = p_txt.add_run(esp)
            run_esp.italic = True
            run_esp.font.name = "Arial"

            run_fam = p_txt.add_run(f" ({fam})")
            run_fam.font.name = "Arial"

            idx += 1

    doc.add_page_break()


# ========= TABLA FINAL MULTI-ORDEN ==========
if especies_unafoto:
    doc.add_heading("Órdenes con una sola foto", level=2)

    filas = math.ceil(len(especies_unafoto) / 2)
    tabla = doc.add_table(rows=filas, cols=2)

    idx = 0
    for r in range(filas):
        for c in range(2):
            if idx >= len(especies_unafoto):
                break

            orden, esp, fam, ruta = especies_unafoto[idx]
            cell = tabla.rows[r].cells[c]

            cell._element.clear()

            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            insertar_imagen(p_img, ruta)

            p_txt = cell.add_paragraph()
            p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_esp = p_txt.add_run(esp)
            run_esp.italic = True
            run_esp.font.name = "Arial"

            run_info = p_txt.add_run(f" ({fam}) – Orden: {orden}")
            run_info.font.name = "Arial"

            idx += 1


# ========= GUARDAR ==========
doc.save(salida_word)
print("\n✔ Word generado en:", salida_word)
